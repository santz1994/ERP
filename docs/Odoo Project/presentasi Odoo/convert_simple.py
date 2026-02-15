"""
Simple and reliable MD to DOCX converter
Preserves formatting properly
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def is_box_char_line(line):
    """Check if line contains ASCII box drawing characters"""
    box_chars = ['┌', '│', '└', '─', '═', '║', '╔', '╚', '╗', '╝', '├', '┤', '┬', '┴', '┼', '╠', '╣', '╦', '╩', '╬']
    return any(c in line for c in box_chars)

def simple_md_to_docx(md_file, docx_file):
    """Simple markdown to docx conversion"""
    
    # Read markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Process line by line
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Code block
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block - add all lines
                if code_lines:
                    p = doc.add_paragraph()
                    p.style.font.name = 'Consolas'
                    p.style.font.size = Pt(9)
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Box drawing lines (treat as code)
        if is_box_char_line(line):
            if not code_lines:
                code_lines = [line]
            else:
                code_lines.append(line)
            # Look ahead to see if more box lines
            if i + 1 < len(lines) and not is_box_char_line(lines[i + 1]):
                # End of box - add it
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(8)
                code_lines = []
            i += 1
            continue
        
        # Headers
        if line.startswith('# ') and not line.startswith('##'):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        elif line.startswith('## ') and not line.startswith('###'):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('### ') and not line.startswith('####'):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue
        
        # Horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('─' * 70)
            run.font.color.rgb = RGBColor(200, 200, 200)
            i += 1
            continue
        
        # Table
        if '|' in line and line.strip().startswith('|'):
            table_lines.append(line)
            # Look ahead for more table lines
            if i + 1 < len(lines) and '|' in lines[i + 1]:
                i += 1
                continue
            else:
                # End of table - process it
                if len(table_lines) > 1:
                    # Parse table
                    rows = []
                    for tline in table_lines:
                        if '---' in tline or '===' in tline:
                            continue
                        cells = [c.strip() for c in tline.split('|')]
                        cells = [c for c in cells if c]
                        if cells:
                            rows.append(cells)
                    
                    if rows:
                        max_cols = max(len(r) for r in rows)
                        table = doc.add_table(rows=len(rows), cols=max_cols)
                        table.style = 'Light Grid Accent 1'
                        
                        for row_idx, row_data in enumerate(rows):
                            for col_idx, cell_data in enumerate(row_data):
                                if col_idx < max_cols:
                                    cell = table.rows[row_idx].cells[col_idx]
                                    cell.text = cell_data
                                    if row_idx == 0:
                                        set_cell_background(cell, 'D9E1F2')
                                        for para in cell.paragraphs:
                                            for run in para.runs:
                                                run.font.bold = True
                
                table_lines = []
                i += 1
                continue
        
        # Lists
        if line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
            continue
        
        # Numbered lists
        if len(line) > 2 and line[0].isdigit() and line[1:3] == '. ':
            doc.add_paragraph(line[3:], style='List Number')
            i += 1
            continue
        
        # Regular paragraph
        if line.strip():
            p = doc.add_paragraph()
            # Handle bold text
            parts = line.split('**')
            for idx, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    if idx % 2 == 1:  # Odd indices are bold
                        run.font.bold = True
        else:
            # Empty line
            doc.add_paragraph()
        
        i += 1
    
    # Save
    doc.save(docx_file)
    print(f"✓ Converted: {os.path.basename(md_file)}")

def main():
    """Convert all markdown files"""
    md_files = [f for f in os.listdir('.') if f.endswith('.md') and f.startswith(('1_', '2_', '3_', '4_', '5_'))]
    
    if not md_files:
        print("No files to convert")
        return
    
    print(f"\nConverting {len(md_files)} files to DOCX...\n")
    
    for md_file in sorted(md_files):
        docx_file = md_file.replace('.md', '.docx')
        try:
            simple_md_to_docx(md_file, docx_file)
        except Exception as e:
            print(f"✗ Error converting {md_file}: {e}")
    
    print(f"\n✓ Conversion complete!")
    print("All files saved with proper formatting")

if __name__ == '__main__':
    main()
