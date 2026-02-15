"""
Convert Markdown to DOCX using markdown2 and python-docx
"""
import os
import sys

try:
    import markdown2
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.style import WD_STYLE_TYPE
except ImportError as e:
    print(f"❌ Error: Missing required package: {e}")
    print("\nPlease install required packages:")
    print("pip install markdown2 python-docx")
    sys.exit(1)

def md_to_docx(md_file, docx_file):
    """Convert a markdown file to docx"""
    
    # Read markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert markdown to HTML
    html = markdown2.markdown(md_content, extras=['tables', 'fenced-code-blocks', 'strike'])
    
    # Create docx
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Split by lines and process
    lines = md_content.split('\n')
    
    for line in lines:
        line = line.rstrip()
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        
        # Horizontal line
        elif line.startswith('---'):
            doc.add_paragraph()
        
        # Code blocks
        elif line.startswith('```'):
            continue
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        
        # Numbered lists
        elif len(line) > 2 and line[0].isdigit() and line[1:3] == '. ':
            doc.add_paragraph(line[3:], style='List Number')
        
        # Tables (simple detection)
        elif '|' in line and not line.startswith('┌'):
            # Skip table processing for simplicity
            doc.add_paragraph(line)
        
        # Regular paragraph (skip ASCII art)
        elif line and not any(c in line for c in ['┌', '│', '└', '─', '═', '║', '╔', '╚']):
            doc.add_paragraph(line)
    
    # Save
    doc.save(docx_file)
    print(f"✅ Converted: {md_file} → {docx_file}")

def main():
    """Convert all .md files in current directory"""
    
    current_dir = os.getcwd()
    md_files = [f for f in os.listdir(current_dir) if f.endswith('.md')]
    
    if not md_files:
        print("❌ No .md files found in current directory")
        return
    
    print(f"Found {len(md_files)} markdown files\n")
    
    for md_file in md_files:
        docx_file = md_file.replace('.md', '.docx')
        try:
            md_to_docx(md_file, docx_file)
        except Exception as e:
            print(f"❌ Error converting {md_file}: {e}")
    
    print(f"\n✅ Conversion complete! Created {len(md_files)} .docx files")

if __name__ == '__main__':
    main()
