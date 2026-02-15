"""
Convert Markdown to DOCX with better formatting
Improved version with better ASCII art handling, tables, and layout
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()

def set_cell_background(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def is_ascii_box_line(line):
    """Check if line is part of ASCII box drawing"""
    box_chars = ['┌', '│', '└', '─', '═', '║', '╔', '╚', '╗', '╝', '├', '┤', '┬', '┴', '┼', '╠', '╣', '╦', '╩', '╬', '▼', '▲', '→', '←', '✅', '⚠️', '❌', '🔴', '🟡', '🟢', '📋', '📅', '🔄']
    return any(char in line for char in box_chars)

def clean_emoji_text(text):
    """Clean text but keep emojis readable"""
    return text

def md_to_docx_improved(md_file, docx_file):
    """Convert markdown to docx with improved formatting"""
    
    # Read markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create document
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # Configure styles
    styles = doc.styles
    
    # Normal style
    style_normal = styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    
    # Code block style (for ASCII art)
    try:
        style_code = styles['Code']
    except KeyError:
        style_code = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    style_code.font.name = 'Consolas'
    style_code.font.size = Pt(8)
    style_code.paragraph_format.left_indent = Inches(0.2)
    style_code.paragraph_format.space_before = Pt(6)
    style_code.paragraph_format.space_after = Pt(6)
    
    # Process content
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    in_ascii_box = False
    ascii_box_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Code block detection
        if line.startswith('```'):
            in_code_block = not in_code_block
            if not in_code_block and ascii_box_lines:
                # End of code block - add collected lines
                p = doc.add_paragraph(style='Code')
                p.add_run('\n'.join(ascii_box_lines))
                ascii_box_lines = []
            i += 1
            continue
        
        if in_code_block:
            ascii_box_lines.append(line)
            i += 1
            continue
        
        # ASCII box detection (outside code blocks)
        if is_ascii_box_line(line):
            if not in_ascii_box:
                in_ascii_box = True
                ascii_box_lines = []
            ascii_box_lines.append(line)
            i += 1
            continue
        else:
            # End of ASCII box
            if in_ascii_box and ascii_box_lines:
                p = doc.add_paragraph(style='Code')
                p.add_run('\n'.join(ascii_box_lines))
                ascii_box_lines = []
                in_ascii_box = False
        
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
            i += 1
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            i += 1
            continue
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            i += 1
            continue
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
            i += 1
            continue
        
        # Horizontal rule
        if line == '---':
            p = doc.add_paragraph()
            p.add_run('_' * 80)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        # Table detection
        if '|' in line and not is_ascii_box_line(line):
            # Try to parse as markdown table
            table_lines = [line]
            j = i + 1
            while j < len(lines) and '|' in lines[j]:
                table_lines.append(lines[j].rstrip())
                j += 1
            
            if len(table_lines) >= 2:
                # Parse table
                rows = []
                for tline in table_lines:
                    if '---' in tline or '===' in tline:
                        continue  # Skip separator line
                    cells = [cell.strip() for cell in tline.split('|')]
                    cells = [c for c in cells if c]  # Remove empty
                    if cells:
                        rows.append(cells)
                
                if rows:
                    # Create table
                    max_cols = max(len(row) for row in rows)
                    table = doc.add_table(rows=len(rows), cols=max_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            
                            # Header row formatting
                            if row_idx == 0:
                                set_cell_background(cell, 'D9E1F2')
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                    
                    i = j
                    continue
        
        # Bullet lists
        if line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
            continue
        
        # Numbered lists
        if len(line) > 2 and line[0].isdigit() and line[1:3] == '. ':
            doc.add_paragraph(line[3:], style='List Number')
            i += 1
            continue
        
        # Bold/italic detection (simple)
        if line:
            p = doc.add_paragraph()
            
            # Check for bold markers
            if '**' in line:
                parts = line.split('**')
                for idx, part in enumerate(parts):
                    run = p.add_run(part)
                    if idx % 2 == 1:  # Odd indices are bold
                        run.font.bold = True
            else:
                p.add_run(line)
        else:
            # Empty line - add spacing
            doc.add_paragraph()
        
        i += 1
    
    # Flush remaining ASCII box
    if ascii_box_lines:
        p = doc.add_paragraph(style='Code')
        p.add_run('\n'.join(ascii_box_lines))
    
    # Save
    doc.save(docx_file)
    print(f"✅ Converted: {os.path.basename(md_file)} → {os.path.basename(docx_file)}")

def main():
    """Convert all .md files with improved formatting"""
    
    current_dir = os.getcwd()
    md_files = [f for f in os.listdir(current_dir) if f.endswith('.md')]
    
    if not md_files:
        print("❌ No .md files found")
        return
    
    print(f"🔄 Converting {len(md_files)} files with improved layout...\n")
    
    success = 0
    for md_file in md_files:
        docx_file = md_file.replace('.md', '.docx')
        try:
            md_to_docx_improved(md_file, docx_file)
            success += 1
        except Exception as e:
            print(f"❌ Error converting {md_file}: {e}")
    
    print(f"\n✅ Successfully converted {success}/{len(md_files)} files")
    print(f"📂 Location: {current_dir}")

if __name__ == '__main__':
    main()
