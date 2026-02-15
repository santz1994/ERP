"""Comprehensive emoji checker for docx files - check paragraphs, tables, and runs"""
from docx import Document
import re
import os

def comprehensive_emoji_check(filename):
    """Check for emojis in all parts of docx"""
    doc = Document(filename)
    
    # Common emojis
    common_emojis = ['✅', '⚠️', '❌', '🔴', '🟡', '🟢', '📋', '📅', '🔄', '⏳', 
                     '▼', '▲', '→', '←', '✓', '⚠', '✗', '☑', '☒']
    
    found_emojis = {}
    total_count = 0
    
    # Check paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        for emoji in common_emojis:
            if emoji in text:
                if emoji not in found_emojis:
                    found_emojis[emoji] = []
                found_emojis[emoji].append(f"Para {i}: {text[:50]}")
                total_count += text.count(emoji)
    
    # Check tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text
                for emoji in common_emojis:
                    if emoji in text:
                        if emoji not in found_emojis:
                            found_emojis[emoji] = []
                        found_emojis[emoji].append(f"Table {table_idx} Row {row_idx} Cell {cell_idx}: {text[:30]}")
                        total_count += text.count(emoji)
    
    # Check using unicode regex
    emoji_pattern = re.compile("["
        "\U0001F600-\U0001F64F"  # emoticons
        "\U0001F300-\U0001F5FF"  # symbols & pictographs
        "\U0001F680-\U0001F6FF"  # transport & map
        "\U0001F1E0-\U0001F1FF"  # flags
        "\U00002600-\U000027BF"  # misc symbols (expanded)
        "\U0001F900-\U0001F9FF"
        "\U0001FA00-\U0001FA6F"
        "]+", flags=re.UNICODE)
    
    all_text = '\n'.join([p.text for p in doc.paragraphs])
    unicode_emojis = emoji_pattern.findall(all_text)
    if unicode_emojis:
        for emoji in set(unicode_emojis):
            if emoji not in found_emojis:
                found_emojis[emoji] = []
            found_emojis[emoji].append(f"Unicode emoji found")
            total_count += all_text.count(emoji)
    
    return found_emojis, total_count

# Check all docx files
docx_files = [f for f in os.listdir('.') if f.endswith('.docx')]

print("\n=== COMPREHENSIVE EMOJI CHECK ===\n")

all_clean = True
for docx_file in docx_files:
    print(f"Checking: {docx_file}")
    emojis, total = comprehensive_emoji_check(docx_file)
    
    if emojis or total > 0:
        print(f"  FOUND {total} emojis:")
        for emoji, locations in emojis.items():
            print(f"    {repr(emoji)}: {len(locations)} occurrences")
            for loc in locations[:3]:  # Show first 3
                print(f"      - {loc}")
        print()
        all_clean = False
    else:
        print(f"  CLEAN (no emojis)")
    print()

if all_clean:
    print("ALL FILES ARE EMOJI-FREE!")
else:
    print("SOME FILES STILL HAVE EMOJIS - NEED TO REPROCESS")
